"""
Append fraud-workflow sections to the existing demo .docx files in
finwatch/docs/demo/. The .docx files were authored manually (no source
generator); this script preserves their existing content and only adds
new paragraphs at the end, matching the observed style convention:

    sz=15.0 bold=True   -> top-level section heading
    sz=12.0 bold=True   -> subsection heading
    sz=10.0 monospace   -> code block (Consolas)
    default             -> prose paragraph

Run once per docs revision. The script is idempotent in the sense that
existing content is never modified; only the new fraud-workflow
appendix is added. Run again to re-add (you'd want to manually remove
any prior appendix first).
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor


DEMO_DIR = Path(__file__).resolve().parent.parent / "docs" / "demo"
DETAILED = DEMO_DIR / "demo_huong_dan_chi_tiet.docx"
SUMMARY = DEMO_DIR / "demo_huong_dan_tom_tat.docx"


# ---------- paragraph helpers ----------

def add_heading(doc, text, *, size=15, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_prose(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    return p


# ---------- detailed appendix ----------

def append_detailed(doc):
    add_heading(doc, "10. Lớp Fraud Workflow (cập nhật mới)", size=15)
    add_prose(doc,
        "Phiên bản này bổ sung lớp fraud-workflow trên nền pipeline CDC: "
        "bảng case log fraud_alerts trong Postgres, worker chạy 6 anomaly "
        "query rồi ghi case về PG, trang quản trị tài khoản với nút khoá/mở "
        "khoá, và hàng đợi case trong UI. Đây là phần biến demo từ 'showcase "
        "pipeline' thành 'hệ thống fraud-ops thực'."
    )

    add_heading(doc, "10.1. Service mới: fraud-worker", size=12)
    add_prose(doc,
        "docker-compose.yml giờ có 9 service. Service thứ 9 là fraud-worker, "
        "chạy scripts/fraud_alert_worker.py --interval 30. Mỗi 30 giây nó "
        "gọi 6 anomaly query trong clickhouse/queries/anomaly_*.sql, tính "
        "severity theo bảng spec (LARGE_AMT > 500M -> critical, ZSCORE |z|>5 "
        "-> critical, các trường hợp khác theo bảng), gộp theo account và "
        "insert dòng mới vào bảng Postgres fraud_alerts. Dedup: 1 giờ theo "
        "cặp (account_id, rule_code)."
    )
    add_code(doc,
        "docker compose up -d fraud-worker\n"
        "docker compose logs -f fraud-worker         # tail log\n"
        "python scripts/fraud_alert_worker.py --once # ad-hoc 1 tick từ host"
    )
    add_prose(doc,
        "Log mẫu mỗi tick (một dòng / rule): "
        "[worker] rule=VELOCITY new=3 dedup=12 skipped=0 elapsed=180ms. "
        "skipped đếm các account_id có trong ClickHouse nhưng đã bị xoá ở "
        "Postgres (FK violation) — bị skip an toàn, không huỷ tick của rule."
    )

    add_heading(doc, "10.2. Bảng fraud_alerts (PG) và CDC qua ClickHouse", size=12)
    add_prose(doc,
        "Schema mới ở postgres/init/02_fraud_workflow.sql: id UUID PK, "
        "account_id UUID FK, rule_code (VELOCITY/LARGE_AMT/MULTI_CCY/ZSCORE/"
        "HIGH_RISK/FAIL_SPIKE), severity (low/medium/high/critical), txn_count "
        "INT, total_amount DECIMAL(18,2), evidence JSONB, status (open/"
        "closed_fraud/closed_clean), created_at, resolved_at. Index dedup "
        "(account_id, rule_code, created_at). Publication finwatch_pub được "
        "ALTER để thêm public.fraud_alerts; connector JSON cũng được cập "
        "nhật table.include.list. Phía ClickHouse: clickhouse/init/"
        "05_fraud_alerts.sql tạo Kafka engine, target ReplacingMergeTree và "
        "MV theo đúng pattern của 3 bảng cũ (decimal-as-string, ISO "
        "timestamp parse, __source_ts_ms làm version column)."
    )
    add_code(doc,
        "docker exec finwatch-postgres psql -U finwatch -d finwatch -c \"\\d fraud_alerts\"\n"
        "docker exec finwatch-clickhouse clickhouse-client -q \\\n"
        "  \"SHOW CREATE TABLE finwatch.fraud_alerts\"\n"
        "docker exec finwatch-postgres psql -U finwatch -d finwatch -c \\\n"
        "  \"SELECT tablename FROM pg_publication_tables WHERE pubname='finwatch_pub'\"\n"
        "# Kỳ vọng: 4 dòng (accounts, merchants, transactions, fraud_alerts)"
    )

    add_heading(doc, "10.3. Balance ledger + lock/unlock — closed loop", size=12)
    add_prose(doc,
        "Mỗi đường insert giao dịch (API /api/insert-transaction, generator, "
        "simulator) giờ chạy trong một transaction Postgres với "
        "SELECT ... FOR UPDATE trên account. Nếu account.status != 'active' "
        "-> txn ghi status='failed' description='rejected: account <status>', "
        "không trừ balance. Nếu type là debit và balance < amount -> txn ghi "
        "status='failed' description='insufficient funds', không trừ balance. "
        "Ngược lại -> txn ghi status='completed' và UPDATE accounts SET "
        "balance = balance +/- amount. Logic này nằm ở tầng application "
        "(không phải trigger DB) để simulator (insert thẳng vào PG bỏ qua "
        "API) vẫn tuân thủ."
    )
    add_prose(doc,
        "API lock/unlock đảo accounts.status: POST /api/accounts/{id}/lock "
        "set 'suspended' (409 nếu không active), POST /api/accounts/{id}/"
        "unlock set 'active' (409 nếu không suspended). Sau khi gọi lock, "
        "trong vòng ~1 giây cả 3 đường insert đều route txn của account này "
        "sang status='failed'."
    )
    add_code(doc,
        "# Closed-loop test (PowerShell)\n"
        "$ACC = (docker exec finwatch-postgres psql -U finwatch -d finwatch -tAc \\\n"
        "  \"SELECT id FROM accounts WHERE email='nguyenvana@email.com'\").Trim()\n"
        "$body = \"{`\"account_id`\":`\"$ACC`\",`\"merchant`\":`\"VinMart`\",\" + `\n"
        "        \"`\"amount`\":100000,`\"type`\":`\"purchase`\"}\"\n"
        "\n"
        "curl.exe -s -X POST \"http://localhost:3002/api/accounts/$ACC/lock\"\n"
        "# {\"status\":\"suspended\"}\n"
        "\n"
        "curl.exe -s -X POST http://localhost:3002/api/insert-transaction \\\n"
        "  -H \"Content-Type: application/json\" -d $body\n"
        "# {\"accepted\":false,\"reason\":\"suspended\",...}\n"
        "\n"
        "curl.exe -s -X POST \"http://localhost:3002/api/accounts/$ACC/unlock\"\n"
        "# {\"status\":\"active\"}\n"
        "\n"
        "curl.exe -s -X POST http://localhost:3002/api/insert-transaction \\\n"
        "  -H \"Content-Type: application/json\" -d $body\n"
        "# {\"accepted\":true,\"new_balance\":<balance - 100000>,...}"
    )

    add_heading(doc, "10.4. Trang /accounts + /accounts/[id]", size=12)
    add_prose(doc,
        "/accounts: danh bạ tài khoản với ô tìm theo tên/email (debounce 300 "
        "ms), cột name | email | balance | status badge | open_alerts_24h | "
        "View. Cột open_alerts_24h join sang fraud_alerts FINAL trong "
        "ClickHouse, đếm case status='open' trong 24h gần nhất. /accounts/"
        "[id]: header có balance lớn, status badge, nút Suspend/Reactivate; "
        "hai panel side-by-side hiển thị 20 transactions gần nhất + 20 alert "
        "history. SWR refresh 5 giây."
    )
    add_prose(doc,
        "Kịch bản 'wow' khi demo: ở Dashboard click card-cloning (15 txn) -> "
        "ở /accounts search Do Van I -> mở View -> bảng Alert history có "
        "VELOCITY high 15 -> click Suspend -> ở Dashboard click card-cloning "
        "lại -> quay về trang chi tiết thấy 15 dòng failed 'rejected: account "
        "suspended' đỏ rực. Click Reactivate, status quay lại active. Toàn "
        "bộ chu trình detect -> action -> reject diễn ra trong ~5 click."
    )

    add_heading(doc, "10.5. Trang /alerts — hàng đợi case persistent", size=12)
    add_prose(doc,
        "/alerts là hàng đợi đọc từ fraud_alerts FINAL với 6 chip filter rule "
        "(VELOCITY/LARGE_AMT/MULTI_CCY/ZSCORE/HIGH_RISK/FAIL_SPIKE) và 4 "
        "chip filter severity. Khác Alert feed ở Dashboard (tính realtime "
        "từ transactions FINAL, transient), trang /alerts là case persistent "
        "do worker viết — analyst xử lý từ đây. Cột Account là link "
        "clickable sang /accounts/[id]: workflow 4 click hoàn chỉnh = lọc "
        "critical -> click account -> click Suspend."
    )

    add_heading(doc, "10.6. Verification cho phần fraud-workflow", size=12)
    add_code(doc,
        "# 1. Bảng + publication\n"
        "docker exec finwatch-postgres psql -U finwatch -d finwatch -c \"\\d fraud_alerts\"\n"
        "docker exec finwatch-postgres psql -U finwatch -d finwatch -c \\\n"
        "  \"SELECT tablename FROM pg_publication_tables WHERE pubname='finwatch_pub'\"\n"
        "\n"
        "# 2. Worker tick\n"
        "python scripts/simulate_fraud.py --scenario card-cloning\n"
        "python scripts/fraud_alert_worker.py --once\n"
        "docker exec finwatch-postgres psql -U finwatch -d finwatch -c \\\n"
        "  \"SELECT rule_code, severity, count(*) FROM fraud_alerts GROUP BY 1, 2\"\n"
        "\n"
        "# 3. Case flow qua CDC (đợi ~10s)\n"
        "docker exec finwatch-clickhouse clickhouse-client -q \\\n"
        "  \"SELECT rule_code, count() FROM finwatch.fraud_alerts FINAL\n"
        "   WHERE cdc_op != 'd' GROUP BY 1\"\n"
        "\n"
        "# 4. UI routes\n"
        "curl.exe -s -o NUL -w \"%{http_code}\\n\" http://localhost:3002/accounts\n"
        "curl.exe -s -o NUL -w \"%{http_code}\\n\" http://localhost:3002/alerts\n"
        "# expect: 200, 200\n"
        "\n"
        "# 5. Pytest không regress\n"
        "python -m pytest -q --ignore=tests/test_stress.py\n"
        "# expect: 22 passed"
    )

    add_heading(doc, "10.7. Tóm tắt file thay đổi", size=12)
    add_prose(doc,
        "Mới: postgres/init/02_fraud_workflow.sql, clickhouse/init/"
        "05_fraud_alerts.sql, scripts/fraud_alert_worker.py, web/app/"
        "accounts/page.tsx + [id]/page.tsx, web/app/alerts/page.tsx, web/"
        "app/api/accounts/[id]/{lock,unlock,route,transactions,alerts}/"
        "route.ts, web/app/api/accounts/search/route.ts, web/app/api/alerts/"
        "route.ts. Sửa: web/app/api/insert-transaction/route.ts (transaction "
        "+ ledger), scripts/generate_transactions.py (ledger), scripts/"
        "simulate_fraud.py (_insert ledger), web/components/NavBar.tsx "
        "(thêm 2 link Accounts + Alerts), docker-compose.yml (service "
        "fraud-worker), debezium/connectors/finwatch-connector.json "
        "(table.include.list), docs/architecture.md (section Fraud workflow "
        "layer + Mermaid sequence), docs/runbook.md (subsection Fraud "
        "worker). Tổng: 5 commit (xem git log)."
    )


# ---------- summary appendix ----------

def append_summary(doc):
    add_heading(doc, "Cập nhật: Lớp Fraud Workflow", size=15)
    add_prose(doc,
        "Bản này thêm closed-loop fraud-ops: bảng case log fraud_alerts trong "
        "Postgres (replicated qua CDC sang ClickHouse), worker chạy 6 "
        "anomaly query mỗi 30 giây và ghi case về PG, hai trang UI mới "
        "/accounts và /alerts, plus nút lock/unlock và balance ledger ở 3 "
        "đường insert."
    )

    add_heading(doc, "Service mới + 2 trang UI mới", size=12)
    add_prose(doc,
        "docker-compose có 9 service (thêm fraud-worker). NavBar có 8 link "
        "(thêm Accounts và Alerts giữa Fraud rules và Kafka)."
    )
    add_code(doc,
        "docker compose up -d fraud-worker\n"
        "python scripts/fraud_alert_worker.py --once  # ad-hoc tick\n"
        "# UI: http://localhost:3002/accounts  http://localhost:3002/alerts"
    )

    add_heading(doc, "Closed-loop test (60 giây)", size=12)
    add_code(doc,
        "# Lock account, insert thử -> reject; unlock, insert lại -> accept\n"
        "$ACC = (docker exec finwatch-postgres psql -U finwatch -d finwatch -tAc \\\n"
        "  \"SELECT id FROM accounts WHERE email='nguyenvana@email.com'\").Trim()\n"
        "curl.exe -s -X POST \"http://localhost:3002/api/accounts/$ACC/lock\"\n"
        "# {\"status\":\"suspended\"}\n"
        "curl.exe -s -X POST http://localhost:3002/api/insert-transaction \\\n"
        "  -H \"Content-Type: application/json\" \\\n"
        "  -d \"{`\"account_id`\":`\"$ACC`\",`\"merchant`\":`\"VinMart`\",\" + `\n"
        "      \"`\"amount`\":100000,`\"type`\":`\"purchase`\"}\"\n"
        "# {\"accepted\":false,\"reason\":\"suspended\",...}\n"
        "curl.exe -s -X POST \"http://localhost:3002/api/accounts/$ACC/unlock\"\n"
        "# {\"status\":\"active\"}"
    )

    add_heading(doc, "Demo flow update — Phần 'wow' 2 phút", size=12)
    add_prose(doc,
        "Sau khi click card-cloning ở Dashboard, mở tab Accounts -> search "
        "tên người Việt -> mở chi tiết -> Suspend -> chạy generator lại -> "
        "thấy txn mới của account này landed là status='failed' với "
        "description 'rejected: account suspended'. Bằng chứng pipeline phản "
        "ứng < 1 giây và logic áp dụng đồng nhất ở cả 3 đường insert."
    )

    add_heading(doc, "Verification tóm tắt", size=12)
    add_code(doc,
        "docker compose ps                       # 9 service Up\n"
        "python scripts/simulate_fraud.py --scenario card-cloning\n"
        "python scripts/fraud_alert_worker.py --once\n"
        "docker exec finwatch-postgres psql -U finwatch -d finwatch -c \\\n"
        "  \"SELECT rule_code, count(*) FROM fraud_alerts GROUP BY 1\"\n"
        "# >=1 dòng (VELOCITY)\n"
        "docker exec finwatch-clickhouse clickhouse-client -q \\\n"
        "  \"SELECT rule_code, count() FROM finwatch.fraud_alerts FINAL\n"
        "   WHERE cdc_op != 'd' GROUP BY 1\"\n"
        "# same"
    )


# ---------- driver ----------

def main():
    for path, appender in [(DETAILED, append_detailed), (SUMMARY, append_summary)]:
        if not path.exists():
            print(f"[skip] {path} not found")
            continue
        doc = Document(str(path))
        appender(doc)
        doc.save(str(path))
        print(f"[ok] appended fraud-workflow section to {path.name}")


if __name__ == "__main__":
    main()
